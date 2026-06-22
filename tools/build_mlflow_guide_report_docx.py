from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / ".opencode" / "skills" / "MLFLOW_5_STEP_GUIDE.docx"

FONT = "Malgun Gothic"
TITLE_BLUE = RGBColor(18, 67, 115)
BLUE = RGBColor(37, 99, 160)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(95, 103, 115)
LIGHT_BLUE = "EAF3FB"
PALE_BLUE = "F5FAFE"
LIGHT_GRAY = "F6F8FA"
GREEN = "EAF7EF"
GOLD = "FFF7E6"
BORDER = "D7E1EC"


def set_font(run, size=10.5, bold=None, color=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = INK
    normal.paragraph_format.line_spacing = 1.14
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color, before, after in [
        ("Heading 1", 16.5, TITLE_BLUE, 15, 7),
        ("Heading 2", 13.2, BLUE, 10, 5),
        ("Heading 3", 11.2, TITLE_BLUE, 7, 3),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("MLflow 5단계 Skill Guide")
    set_font(r, size=8.5, color=MUTED)


def para(doc, text="", size=10.2, bold=False, color=INK, after=5, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.14
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, size=10.0, color=INK)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, size=10.0, color=INK)


def code_box(doc, lines, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    table_borders(table, color="D8DEE8", size="4")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell_margins(cell, top=120, bottom=120, start=150, end=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(lines):
        if index:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.size = Pt(8.8)
        r.font.color.rgb = RGBColor(30, 41, 59)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def two_col_table(doc, rows, header=("구분", "내용"), col_widths=(1.65, 4.85)):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for i, width in enumerate(col_widths):
        table.columns[i].width = Inches(width)
    table_borders(table)
    repeat_header(table.rows[0])
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = text
        shade_cell(cell, LIGHT_BLUE)
        cell_margins(cell)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_font(r, size=9.2, bold=row_index == 0, color=TITLE_BLUE if row_index == 0 else INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)
    table_borders(table)
    repeat_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        shade_cell(cell, LIGHT_BLUE)
        cell_margins(cell)
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = value
            cell_margins(cells[i])
            if i == 0 and value.isdigit():
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_font(r, size=9.0, bold=row_index == 0, color=TITLE_BLUE if row_index == 0 else INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def callout(doc, title, body, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    table_borders(table, color="C8D8EA")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell_margins(cell, top=140, bottom=140, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title)
    set_font(r, size=10.2, bold=True, color=TITLE_BLUE)
    p.add_run("\n")
    r2 = p.add_run(body)
    set_font(r2, size=9.6, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def section_break(doc, title, subtitle):
    doc.add_section(WD_SECTION.NEW_PAGE)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title)
    set_font(r, size=20, bold=True, color=TITLE_BLUE)
    para(doc, subtitle, size=10.5, color=MUTED, after=16)


def build():
    doc = Document()
    configure_doc(doc)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(44)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("MLflow 5단계 Skill 기술 아키텍처 및 운영 설계서")
    set_font(r, size=27, bold=True, color=TITLE_BLUE)
    para(doc, "AI Studio 모델 프로젝트 온보딩 · 검증 · MLflow 기록 체계", size=12, color=MUTED, after=18)
    callout(
        doc,
        "Executive Summary",
        "본 문서는 ML 개발자가 사용자가 지정한 모델 프로젝트 폴더를 기준으로 프로젝트 분석, 환경 검증, 로컬 모델 생성, 추론 테스트, MLflow 기록 확인을 챗봇 기반으로 수행하기 위한 5단계 Skill 운영 아키텍처를 정의한다. 핵심은 모델 프로젝트 계약을 표준화하고, 실행 전 검증과 산출물 확인을 자동화해 AI Studio에서 재사용 가능한 모델 온보딩 흐름을 제공하는 것이다.",
        fill=PALE_BLUE,
    )
    two_col_table(
        doc,
        [
            ("문서 목적", "MLflow 5단계 Skill의 실행 계약, 운영 흐름, AI Studio 전환 기준을 정의"),
            ("주요 사용자", "PM, AI Platform Architect, MLOps 담당자, ML 개발 리더"),
            ("대상 범위", "사용자가 지정한 모델 프로젝트 폴더"),
            ("필수 계약", "aiu_custom/ · local_serving/ · save_model/ · ai_studio.env"),
            ("주요 산출물", "save_model/ 모델 산출물, inference 결과, MLflow Run/Artifact/Model Registry 기록"),
            ("보안 원칙", "secret 값과 mlflow_tracking_password 값은 출력하지 않고 상태만 표시"),
        ],
    )
    matrix_table(
        doc,
        ["단계", "검증 영역", "핵심 책임"],
        [
            ("1", "Project Analyze", "모델 존재 여부, entrypoint, 필수 폴더, 샘플 선택 여부 판단"),
            ("2", "Environment Check", "Python, dependency, MLflow, ai_studio.env 설정 검증"),
            ("3", "Train Model", "기존 프로젝트 실행 또는 표준 샘플 기반 모델 산출물 생성"),
            ("4", "Inference Test", "input_example 기반 로드 및 predict 가능 여부 검증"),
            ("5", "MLflow Verify", "Run, artifact, pyfunc model, Model Registry 기록 확인"),
        ],
        [0.55, 1.65, 4.3],
    )

    section_break(doc, "1. 설계 배경과 범위", "챗봇 기반 모델 온보딩을 표준화하기 위한 판단 기준을 정의한다.")
    doc.add_heading("1.1 문제 정의", level=1)
    para(
        doc,
        "ML 개발자가 로컬 모델 프로젝트를 AI Studio와 MLflow로 연결할 때 가장 큰 비용은 프로젝트마다 다른 구조, 불명확한 실행 환경, 모델 산출물 위치, 추론 진입점, MLflow 등록 상태를 매번 수동으로 확인해야 한다는 점이다. 5단계 Skill은 이 과정을 대화형 자동 점검 흐름으로 바꾸는 운영 계층이다.",
        after=8,
    )
    doc.add_heading("1.2 설계 목표", level=2)
    for item in [
        "모델 프로젝트 폴더를 기준으로 실행 가능한 모델이 있는지 일관되게 판단한다.",
        "학습 전 필수 설정과 디렉터리 계약을 검증해 실패를 앞 단계에서 차단한다.",
        "모델 생성, 추론 테스트, MLflow 기록 확인을 하나의 연속된 검증 흐름으로 연결한다.",
        "AI Studio 화면으로 이전 가능한 API/상태 모델을 확보한다.",
    ]:
        bullet(doc, item)

    doc.add_heading("1.3 범위와 비범위", level=2)
    matrix_table(
        doc,
        ["구분", "포함", "제외"],
        [
            ("프로젝트 분석", "사용자 지정 모델 프로젝트 폴더 구조 분석", "임의 외부 저장소 자동 다운로드"),
            ("모델 실행", "기존 모델 실행 또는 표준 샘플 기반 모델 생성", "표준 샘플 외 임의 샘플 자동 선택"),
            ("관측성", "MLflow Run, artifact, registered model 확인", "MLflow UI 자체 재구현"),
            ("보안", "secret 비노출, 설정 상태만 표시", "credential 발급 및 권한 정책 관리"),
        ],
        [1.2, 2.7, 2.6],
    )

    section_break(doc, "2. 프로젝트 계약", "모델 프로젝트 폴더가 만족해야 하는 최소 구조와 설정을 정의한다.")
    doc.add_heading("2.1 상위 구조", level=1)
    code_box(
        doc,
        [
            "<model-project-folder>/",
            "  ai_studio.env",
            "  aiu_custom/",
            "  local_serving/",
            "  save_model/",
            "  input_example.json",
            "  train.py 또는 run_model.py",
        ],
    )
    doc.add_heading("2.2 필수/선택 구성", level=2)
    two_col_table(
        doc,
        [
            ("필수 폴더", "aiu_custom/ · local_serving/ · save_model/"),
            ("필수 설정", "ai_studio.env"),
            ("권장 파일", "input_example.json · train.py · run_model.py"),
            ("필수 아님", "offline_weather_agent_core/ · registry/"),
        ],
    )
    doc.add_heading("2.3 ai_studio.env 계약", level=2)
    code_box(
        doc,
        [
            'mlflow_tracking_url=""',
            'mlflow_tracking_username=""',
            'mlflow_tracking_password=""',
            'mlflow_experiment_name=""',
            'mlflow_register_model_name=""',
        ],
    )
    callout(
        doc,
        "Credential Handling",
        "검증 결과에는 set / empty / missing 상태만 표시한다. mlflow_tracking_password를 포함한 secret 값은 로그, 응답, trace, 문서화된 결과에 출력하지 않는다.",
        fill=GOLD,
    )

    doc.add_heading("2.4 모델 존재 여부 판단", level=2)
    code_box(
        doc,
        [
            "모델이 있음",
            "  -> 사용자가 지정한 모델 프로젝트 폴더를 분석해서 그대로 실행한다.",
            "",
            "모델이 없음",
            "  -> .opencode/samples 아래 표준 샘플 3개 중 하나를 자동 선택한다.",
            "  -> 선택한 샘플을 작업 경로에 준비해서 모델을 생성하고 테스트한다.",
        ],
    )
    doc.add_heading("2.5 허용 표준 샘플", level=2)
    numbered(doc, "sklearn_sample")
    numbered(doc, "pytorch_sample")
    numbered(doc, "tensorflow_sample")
    callout(doc, "분기 원칙", "기존 모델 프로젝트가 확인되면 표준 샘플을 사용하지 않는다. 모델 근거가 없을 때만 위 3개 샘플 중 하나를 순서대로 선택한다.", fill=PALE_BLUE)

    section_break(doc, "3. 5단계 운영 흐름", "각 단계의 목적, 입력, 출력, 실패 기준을 명확히 분리한다.")
    flow_rows = [
        ("1", "Project Analyze", "validate_mlflow_project.py", "모델 있음/없음 판단, 필수 폴더 확인, 샘플 선택"),
        ("2", "Environment Check", "check_environment.py", "Python, dependency, MLflow, ai_studio.env 검증"),
        ("3", "Train Model", "run_training.py", "학습/export 실행, save_model/ 산출물 확인"),
        ("4", "Inference Test", "test_inference.py", "input_example 기반 predict 검증"),
        ("5", "MLflow Verify", "verify_mlflow.py", "Run, artifact, registered model 확인"),
    ]
    matrix_table(doc, ["단계", "이름", "스크립트", "출력"], flow_rows, [0.45, 1.35, 1.85, 2.85])

    doc.add_heading("2.1 데이터 흐름", level=1)
    code_box(
        doc,
        [
            "사용자 지정 모델 프로젝트 폴더",
            "  -> Step 1 구조 분석",
            "  -> Step 2 환경 검증",
            "  -> Step 3 학습 또는 export",
            "  -> save_model/ 산출물 생성",
            "  -> Step 4 추론 테스트",
            "  -> Step 5 MLflow 기록 확인",
        ],
        fill="F8FBFF",
    )

    section_break(doc, "4. 단계별 기술 설계", "챗봇 Skill과 로컬 스크립트가 각 단계에서 수행할 책임을 정의한다.")
    doc.add_heading("4.1 Step 1 · 프로젝트 구조 분석", level=1)
    two_col_table(
        doc,
        [
            ("Skill", "agent-mlflow-skill-project-analyze"),
            ("확인 항목", "train.py, run_model.py, input_example.json, ai_studio.env, aiu_custom/, local_serving/, save_model/"),
            ("모델 있음", "기존 프로젝트를 실행 대상으로 지정하고 samples는 보지 않음"),
            ("모델 없음", "sklearn_sample → pytorch_sample → tensorflow_sample 순서로 선택"),
            ("주요 실패", "sample_not_found, missing_required_dir"),
        ],
    )
    doc.add_heading("4.2 Step 2 · 실행 환경 검증", level=1)
    two_col_table(
        doc,
        [
            ("Skill", "agent-mlflow-skill-environment-check"),
            ("확인 항목", "Python, venv/conda, dependency, MLflow version"),
            ("필수 설정", "ai_studio.env 5개 키"),
            ("출력 방식", "secret 값 비노출, 상태만 set/empty/missing"),
            ("주요 실패", "missing_dependency, missing_env, tracking_unreachable"),
        ],
    )
    doc.add_heading("4.3 Step 3 · 로컬 학습 및 모델 생성", level=1)
    two_col_table(
        doc,
        [
            ("Skill", "agent-mlflow-skill-train-model"),
            ("기존 모델", "selected_project_path 기준으로 train.py 또는 run_model.py 실행"),
            ("샘플 모델", "work/<sample-name> 작업 경로에 샘플 준비 후 실행"),
            ("필수 산출물", "save_model/ 모델 산출물"),
            ("실행 안전", "실제 실행은 명시 요청 또는 --execute 기준"),
        ],
    )
    doc.add_heading("4.4 Step 4 · 추론 테스트", level=1)
    two_col_table(
        doc,
        [
            ("Skill", "agent-mlflow-skill-inference-test"),
            ("추론 경로", "aiu_custom/model_wrapper.py, aiu_custom/predict.py, local_serving/, predict.py"),
            ("모델 경로", "save_model/"),
            ("입력", "input_example.json"),
            ("검증", "predict 실행, 응답 schema, JSON 직렬화 가능 여부"),
        ],
    )
    doc.add_heading("4.5 Step 5 · MLflow Run 및 Model 확인", level=1)
    two_col_table(
        doc,
        [
            ("Skill", "agent-mlflow-skill-mlflow-verify"),
            ("확인 항목", "tracking target, experiment, run, artifact, pyfunc model, registered model"),
            ("출력", "최근 run, artifact 상태, model version, MLflow UI 위치"),
            ("주요 실패", "experiment_missing, run_missing, artifact_missing, registry_missing"),
        ],
    )

    section_break(doc, "5. 운영 통제와 리스크", "실패 원인을 일관된 코드로 분류하고 실행 안전성을 보장한다.")
    doc.add_heading("5.1 실패 분류", level=1)
    code_box(
        doc,
        [
            "missing_required_dir:<name>",
            "missing_env_file:ai_studio.env",
            "missing_env:<key>",
            "missing_dependency",
            "missing_train_entrypoint",
            "sample_not_found",
            "sample_prepare_error",
            "artifact_not_created",
            "artifact_invalid",
            "missing_inference_entrypoint",
            "missing_input_example",
            "model_load_error",
            "predict_error",
            "schema_error",
            "tracking_unreachable",
            "experiment_missing",
            "run_missing",
            "artifact_missing",
            "registry_missing",
            "permission_error",
        ],
    )
    doc.add_heading("5.2 실행 안전 정책", level=1)
    for item in [
        "사용자가 지정한 모델 프로젝트 폴더를 먼저 분석한다.",
        "모델이 있으면 샘플을 사용하지 않는다.",
        "모델이 없을 때만 표준 샘플 3개 중 하나를 선택한다.",
        "샘플 원본은 직접 수정하지 않는다.",
        "실제 학습/추론 실행은 명시 요청 또는 --execute가 있을 때만 수행한다.",
        "secret 값과 mlflow_tracking_password 값은 출력하지 않는다.",
        "외부 다운로드나 원격 등록은 기본 동작으로 가정하지 않는다.",
    ]:
        bullet(doc, item)

    doc.add_heading("5.3 기술 평가 관점", level=1)
    matrix_table(
        doc,
        ["평가 항목", "판단 기준", "통제 방법"],
        [
            ("재현성", "동일 프로젝트에서 동일한 분석/실행 결과를 얻는가", "필수 폴더, env 키, entrypoint를 구조적으로 검증"),
            ("운영성", "실패 원인을 사용자와 운영자가 같은 코드로 이해하는가", "공통 실패 코드와 단계별 출력 계약 사용"),
            ("보안성", "credential이 노출되지 않는가", "값 출력 금지, 상태만 출력, password masking"),
            ("확장성", "AI Studio 화면/API로 옮길 수 있는가", "단계별 상태 모델과 스크립트 책임 분리"),
        ],
        [1.3, 2.65, 2.55],
    )

    section_break(doc, "6. AI Studio 전환 설계", "5단계 흐름은 AI Studio의 모델 온보딩 화면과 API로 이전할 수 있다.")
    matrix_table(
        doc,
        ["AI Studio 화면", "연결되는 단계", "주요 기능"],
        [
            ("모델 프로젝트 선택", "Step 1", "사용자가 모델 프로젝트 폴더 지정"),
            ("프로젝트 분석", "Step 1", "필수 폴더, ai_studio.env, entrypoint 확인"),
            ("환경 검증", "Step 2", "Python, dependency, MLflow 설정 확인"),
            ("학습 실행", "Step 3", "train/export 실행 및 save_model 산출물 확인"),
            ("추론 테스트", "Step 4", "input_example 기반 predict 검증"),
            ("MLflow 기록", "Step 5", "run, artifact, registered model 확인"),
        ],
        [1.75, 1.1, 3.65],
    )
    doc.add_heading("6.1 권장 API 상태 모델", level=1)
    matrix_table(
        doc,
        ["상태", "설명", "대표 필드"],
        [
            ("project_analyzed", "프로젝트 구조 분석 완료", "model_found, framework, entrypoint, required_dirs"),
            ("environment_checked", "실행 환경 검증 완료", "python_version, mlflow_version, env_status"),
            ("model_trained", "학습 또는 export 완료", "model_path, artifact_status, training_summary"),
            ("inference_verified", "추론 테스트 완료", "input_example, output_schema, predict_status"),
            ("mlflow_verified", "MLflow 기록 확인 완료", "experiment, run_id, artifact_uri, model_version"),
        ],
        [1.55, 2.05, 2.9],
    )
    doc.add_heading("6.2 최종 확인 질문", level=1)
    for item in [
        "이 모델 프로젝트 폴더는 어떤 ML 프로젝트인가?",
        "필수 폴더 aiu_custom/local_serving/save_model이 있는가?",
        "ai_studio.env 필수 키가 준비되었는가?",
        "현재 환경에서 실행 가능한가?",
        "학습 또는 export 후 save_model/에 모델이 생성되는가?",
        "생성된 모델은 실제로 추론 가능하고 MLflow에 run/model/artifact 기록이 남는가?",
    ]:
        bullet(doc, item)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
