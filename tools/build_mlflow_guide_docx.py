from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / ".opencode" / "skills" / "MLFLOW_5_STEP_GUIDE.md"
OUTPUT = ROOT / ".opencode" / "skills" / "MLFLOW_5_STEP_GUIDE.docx"

FONT = "Malgun Gothic"
ACCENT = RGBColor(31, 78, 121)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F4F6F8"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 7),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 11.5, ACCENT, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_code_paragraph(doc: Document, lines: list[str]):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.25)
    set_table_borders(table, color="D6DBDF", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=120, bottom=120, start=140, end=140)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(8.8)
        run.font.color.rgb = RGBColor(31, 31, 31)
    doc.add_paragraph()


def add_metadata_table(doc: Document):
    rows = [
        ("대상", "사용자가 지정한 모델 프로젝트 폴더"),
        ("필수 폴더", "aiu_custom/ · local_serving/ · save_model/"),
        ("필수 설정", "ai_studio.env"),
        ("표준 샘플", "sklearn_sample → pytorch_sample → tensorflow_sample"),
        ("보안 원칙", "secret 값과 mlflow_tracking_password 값은 출력하지 않음"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(1.65)
    table.columns[1].width = Inches(4.85)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    table.rows[0].cells[0].text = "구분"
    table.rows[0].cells[1].text = "내용"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=9.5, bold=True, color=ACCENT)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for cell in cells:
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=9.2)
    doc.add_paragraph()


def add_flow_table(doc: Document):
    steps = [
        ("1", "Project Analyze", "모델 프로젝트 폴더를 분석하고 모델 있음/없음 분기를 결정"),
        ("2", "Environment Check", "Python, dependency, MLflow, ai_studio.env 상태 확인"),
        ("3", "Train Model", "기존 모델 실행 또는 표준 샘플 기반 모델 생성"),
        ("4", "Inference Test", "input_example 기반 모델 로드 및 predict 검증"),
        ("5", "MLflow Verify", "Run, artifact, pyfunc model, Model Registry 기록 확인"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [0.55, 1.65, 4.3]
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    set_table_borders(table)
    headers = ["단계", "Skill", "책임"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
    set_repeat_table_header(table.rows[0])
    for number, skill, detail in steps:
        cells = table.add_row().cells
        values = [number, skill, detail]
        for idx, value in enumerate(values):
            cells[idx].text = value
            set_cell_margins(cells[idx])
            if idx == 0:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=9.2, bold=row is table.rows[0], color=ACCENT if row is table.rows[0] else None)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    set_table_borders(table, color="C9D8E8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F7FC")
    set_cell_margins(cell, top=130, bottom=130, start=150, end=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10, bold=True, color=ACCENT)
    p.add_run("\n")
    r2 = p.add_run(body)
    set_run_font(r2, size=9.4)
    doc.add_paragraph()


def add_parsed_markdown(doc: Document, text: str):
    in_code = False
    code_lines: list[str] = []
    skip_title = True

    for raw in text.splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                add_code_paragraph(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            continue

        if skip_title and line.startswith("# "):
            skip_title = False
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif re.match(r"^\d+\.\s+", line):
            paragraph = doc.add_paragraph(style="List Number")
            run = paragraph.add_run(re.sub(r"^\d+\.\s+", "", line))
            set_run_font(run)
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(line[2:].strip())
            set_run_font(run)
        else:
            paragraph = doc.add_paragraph()
            # Preserve inline-code readability without overbuilding markdown spans.
            parts = re.split(r"(`[^`]+`)", line)
            for part in parts:
                if part.startswith("`") and part.endswith("`"):
                    run = paragraph.add_run(part[1:-1])
                    run.font.name = "Consolas"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
                    run.font.size = Pt(9.4)
                    run.font.color.rgb = ACCENT
                else:
                    run = paragraph.add_run(part)
                    set_run_font(run)

    if in_code and code_lines:
        add_code_paragraph(doc, code_lines)


def build_docx():
    source_text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("MLflow 5단계 Skill 설명서")
    set_run_font(run, size=23, bold=True, color=ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("OpenCode 기반 ML 모델 프로젝트 분석, 학습, 추론, MLflow 기록 확인 운영 가이드")
    set_run_font(run, size=10.5, color=MUTED)

    add_callout(
        doc,
        "문서 목적",
        "사용자가 지정한 모델 프로젝트 폴더를 기준으로 모델 존재 여부를 판단하고, 환경 검증부터 MLflow 기록 확인까지 5단계로 수행하는 운영 기준을 정리한다.",
    )
    add_metadata_table(doc)
    add_flow_table(doc)
    doc.add_page_break()
    add_parsed_markdown(doc, source_text)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("MLflow 5단계 Skill Guide")
    set_run_font(r, size=8.5, color=MUTED)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
